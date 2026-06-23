"""Generate PES token supply breakdown report as DOCX."""
from __future__ import annotations

import json
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / f"pes-supply-breakdown-{datetime.now(timezone.utc).strftime('%Y-%m-%d')}.docx"

COLORS = {
    "navy": "0B1020",
    "muted": "6B7280",
    "line": "D9DEE8",
    "light": "F7F8FA",
}


def rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.strip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_run_font(run, size=10, color=None, bold=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, COLORS["navy"])
        set_cell_text(cell, header, bold=True, color="FFFFFF", size=9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    set_run_font(run, size=sizes.get(level, 10), bold=True, color=COLORS["navy"])
    p.paragraph_format.space_before = Pt(10 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(6)


def add_para(doc, text, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, color=color or COLORS["muted"])
    p.paragraph_format.space_after = Pt(4)


def fetch_chain_data() -> dict:
    script = r"""
const { ethers } = require('ethers');
const provider = new ethers.JsonRpcProvider('https://bsc-rpc.publicnode.com');
const PES = '0x40D51d93e3Eb057b3558DA71C7CCdEAa27713E41';
const PRESALE = '0x38882c608F64a8dAA5fbAB9a0712361D72866B6B';
const ZERO = '0x0000000000000000000000000000000000000000';
const DEAD = '0x000000000000000000000000000000000000dEaD';
const MAX = 21_000_000n * 10n**18n;
const erc20 = ['function totalSupply() view returns (uint256)','function balanceOf(address) view returns (uint256)','function name() view returns (string)','function symbol() view returns (string)'];
const presaleAbi = ['function totalTokensAllocated() view returns (uint256)','function totalTokensClaimed() view returns (uint256)','function unclaimedAllocatedTokens() view returns (uint256)','function totalPackagesAllocated() view returns (uint256)'];
const fmt = (v) => ethers.formatUnits(v, 18);
(async () => {
  const pes = new ethers.Contract(PES, erc20, provider);
  const presale = new ethers.Contract(PRESALE, presaleAbi, provider);
  const blockNumber = await provider.getBlockNumber();
  const block = await provider.getBlock(blockNumber);
  const [name, symbol, totalSupply, presaleBal, zeroBal, deadBal, allocated, claimed, unclaimed, packages] = await Promise.all([
    pes.name(), pes.symbol(), pes.totalSupply(),
    pes.balanceOf(PRESALE), pes.balanceOf(ZERO), pes.balanceOf(DEAD),
    presale.totalTokensAllocated(), presale.totalTokensClaimed(), presale.unclaimedAllocatedTokens(), presale.totalPackagesAllocated(),
  ]);
  const burnedSupplyDelta = MAX - totalSupply;
  const other = totalSupply - deadBal - zeroBal - presaleBal - claimed;
  console.log(JSON.stringify({
    blockNumber,
    blockTimeUtc: new Date(Number(block.timestamp) * 1000).toISOString(),
    pesToken: PES,
    tokenName: name,
    tokenSymbol: symbol,
    maxMint: fmt(MAX),
    totalSupply: fmt(totalSupply),
    burnedSupplyDelta: fmt(burnedSupplyDelta),
    zeroBalance: fmt(zeroBal),
    deadBalance: fmt(deadBal),
    presale: PRESALE,
    presaleBalance: fmt(presaleBal),
    allocated: fmt(allocated),
    claimed: fmt(claimed),
    unclaimed: fmt(unclaimed),
    packages: packages.toString(),
    otherBalances: fmt(other),
    claimPct: ((Number(claimed) / Number(allocated)) * 100).toFixed(2),
    deadPct: ((Number(deadBal) / Number(MAX)) * 100).toFixed(1),
  }));
})();
"""
    result = subprocess.run(
        ["node", "-e", script],
        cwd=ROOT,
        capture_output=True,
        text=True,
        check=True,
    )
    return json.loads(result.stdout.strip())


def fmt_num(value: str) -> str:
    try:
        num = float(value.replace(",", ""))
        if num >= 1:
            return f"{num:,.2f}".rstrip("0").rstrip(".")
        return value
    except ValueError:
        return value


def build_document(data: dict) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PES 代币供应与归属统计报告")
    set_run_font(run, size=18, bold=True, color=COLORS["navy"])

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"BSC 主网 · 数据区块 #{data['blockNumber']} · {data['blockTimeUtc'][:19].replace('T', ' ')} UTC"
    )
    set_run_font(run, size=10, color=COLORS["muted"])

    add_para(doc, f"代币合约：{data['pesToken']}")
    add_para(doc, f"代币名称：{data['tokenName']} ({data['tokenSymbol']})")

    add_heading(doc, "一、已产出总量", 2)
    add_table(
        doc,
        ["项目", "数量 (PES)", "说明"],
        [
            ["初始铸造（最大供应）", fmt_num(data["maxMint"]), "部署时一次性铸造"],
            ["当前 totalSupply", fmt_num(data["totalSupply"]), "链上总供应量，当前未减少"],
        ],
        widths=[2.0, 1.6, 2.8],
    )

    add_heading(doc, "二、用户购买到钱包的总量", 2)
    add_para(doc, "预售/归属采用「先登记、再按周期释放到钱包」，以下以已领取到账为准。")
    add_table(
        doc,
        ["项目", "数量 (PES)"],
        [
            ["归属登记总量", fmt_num(data["allocated"])],
            ["已发到用户钱包", fmt_num(data["claimed"])],
            ["尚未到账", fmt_num(data["unclaimed"])],
        ],
        widths=[2.5, 2.5],
    )
    add_para(
        doc,
        f"说明：归属合约共 {data['packages']} 份套餐，每份 3,000 PES；已到账比例约 {data['claimPct']}%。"
        f"用户钱包实际持有约 {fmt_num(data['claimed'])} PES（不含二级市场买入）。",
    )

    add_heading(doc, "三、分币（归属）合约剩余量", 2)
    add_table(
        doc,
        ["项目", "数量 (PES)", "说明"],
        [
            ["归属合约地址", data["presale"], ""],
            ["链上 PES 余额", fmt_num(data["presaleBalance"]), "合约内可发放余额"],
            ["账面待释放", fmt_num(data["unclaimed"]), "已登记未领取"],
        ],
        widths=[1.6, 1.6, 2.8],
    )
    shortfall = float(data["unclaimed"]) - float(data["presaleBalance"])
    add_para(
        doc,
        f"要点：账面还需释放 {fmt_num(data['unclaimed'])} PES，但合约内仅剩 {fmt_num(data['presaleBalance'])} PES，"
        f"缺口约 {fmt_num(str(shortfall))} PES，需补款后才能继续发放。",
        color=COLORS["navy"],
    )

    add_heading(doc, "四、销毁量", 2)
    add_table(
        doc,
        ["统计方式", "数量 (PES)", "说明"],
        [
            ["供应量减少（21M − totalSupply）", fmt_num(data["burnedSupplyDelta"]), "未通过减供应方式销毁"],
            ["0x000...0000 地址余额", fmt_num(data["zeroBalance"]), "转入零地址"],
            [
                "0x000...dEaD 地址余额",
                fmt_num(data["deadBalance"]),
                "转入黑洞地址，未减少 totalSupply",
            ],
        ],
        widths=[2.0, 1.6, 2.8],
    )
    add_para(
        doc,
        f"该 PES 合约将代币转入 dead 地址时不会减少 totalSupply。链上官方销毁（减供应）为 0；"
        f"黑洞地址锁定约 {fmt_num(data['deadBalance'])} PES（约总量 {data['deadPct']}%），实质不可流通。",
    )

    add_heading(doc, "五、21,000,000 PES 大致去向（对账）", 2)
    add_table(
        doc,
        ["去向", "约计 (PES)"],
        [
            ["黑洞地址（dead）", fmt_num(data["deadBalance"])],
            ["用户钱包（已领取）", fmt_num(data["claimed"])],
            ["零地址", fmt_num(data["zeroBalance"])],
            ["归属合约余额", fmt_num(data["presaleBalance"])],
            ["其他地址（LP、流通等）", fmt_num(data["otherBalances"])],
            ["合计", fmt_num(data["totalSupply"])],
        ],
        widths=[3.2, 2.0],
    )

    add_heading(doc, "六、关键合约地址", 2)
    add_table(
        doc,
        ["角色", "地址"],
        [
            ["PES Token", data["pesToken"]],
            ["归属合约", data["presale"]],
            ["Presale Owner", "0x042De5f683cA15633A208d2D23dB4F7F47A08821"],
            ["Keeper", "0x7123A25d205190e6844712Cb18e39d6DD5316143"],
        ],
        widths=[1.8, 4.4],
    )

    add_para(doc, f"报告生成时间：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')} UTC")
    add_para(doc, f"文件路径：{OUTPUT.as_posix()}", size=9)

    return doc


def main() -> None:
    try:
        data = fetch_chain_data()
    except Exception as error:
        print(f"链上数据获取失败: {error}", file=sys.stderr)
        sys.exit(1)

    doc = build_document(data)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(str(OUTPUT))


if __name__ == "__main__":
    main()

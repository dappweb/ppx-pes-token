from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("output/DAPPWEB-PES-Contract-Audit-Report-2026-05-26.docx")

COLORS = {
    "navy": "0B1020",
    "gold": "D8A31A",
    "muted": "6B7280",
    "line": "D9DEE8",
    "red": "B42318",
    "orange": "B54708",
    "green": "0F8A5F",
    "light": "F7F8FA",
}


def rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.strip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_run_font(run, size=10, color=None, bold=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color="D9DEE8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, COLORS["navy"])
        set_cell_text(cell, header, bold=True, color="FFFFFF", size=8.5)
        set_cell_borders(cell, "FFFFFF")
        if widths:
            cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=8.2)
            set_cell_borders(cells[idx], COLORS["line"])
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=15 if level == 1 else 12, color=COLORS["navy"], bold=True)
    return p


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=9.5)
    return p


def add_finding(doc, code, severity, title, evidence, impact, recommendation, status="Open"):
    severity_color = {
        "Critical": COLORS["red"],
        "High": COLORS["red"],
        "Medium": COLORS["orange"],
        "Low": COLORS["gold"],
        "Info": COLORS["muted"],
    }.get(severity, COLORS["muted"])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{code}  [{severity}] {title}")
    set_run_font(run, size=12, color=severity_color, bold=True)
    add_para(doc, f"状态：{status}", "状态：")
    add_para(doc, f"证据：{evidence}", "证据：")
    add_para(doc, f"影响：{impact}", "影响：")
    add_para(doc, f"建议：{recommendation}", "建议：")


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    for style_name in ["Normal", "Body Text", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = rgb(COLORS["navy"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DW  DAPPWEB")
    set_run_font(run, size=18, color=COLORS["gold"], bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PES Token & Presale Vesting\n智能合约安全审计报告")
    set_run_font(run, size=24, color=COLORS["navy"], bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Web3 Security · Smart Contract Audit · Launch Readiness")
    set_run_font(run, size=11, color=COLORS["muted"])

    add_table(
        doc,
        ["项目", "内容"],
        [
            ("审计方", "DAPPWEB"),
            ("项目仓库", "ppx-pes-token"),
            ("审计日期", "2026-05-26（Asia/Shanghai）"),
            ("目标网络", "BSC Mainnet / chainId 56"),
            ("报告版本", "v1.1（乱码修复版）"),
        ],
        widths=[1.6, 4.6],
    )
    add_para(doc, "品牌说明：本报告采用 DAPPWEB 的 Web3 安全与产品工程交付口径，重点覆盖合约权限、升级安全、业务逻辑、测试证据、部署状态和上线操作风险。")
    add_para(doc, "免责声明：本报告基于当前仓库代码、自动化测试结果和只读链上状态核验，不构成对未来代码变更、私钥安全、第三方服务或市场表现的保证。")

    doc.add_page_break()

    add_heading(doc, "1. 执行摘要", 1)
    add_para(doc, "DAPPWEB 对 PES Token 与 PESPresaleVesting 合约进行了面向上线风险的安全审计。审计覆盖 Solidity 源码、UUPS 升级路径、预售/释放业务逻辑、测试用例和 BSC 主网部署状态。")
    add_para(doc, "总体结论：未发现 Critical 或 High 等级的直接可利用漏洞。当前主要风险为中心化 Owner 权限较大、释放进度由 Owner 手动推进、以及预售/释放资金保障依赖运营流程。")
    add_para(doc, "上线判断：当前合约可以继续用于受控预售，但正式扩大销售和开放领取前，应完成多签/权限硬化、释放规则确认、链上参数冻结或公告、以及领取前余额复核。")
    add_table(
        doc,
        ["严重级别", "数量", "说明"],
        [
            ("Critical", "0", "未发现可直接导致资金立即被盗或合约完全失控的问题"),
            ("High", "0", "未发现高概率高影响的可利用漏洞"),
            ("Medium", "2", "Owner 权限集中、释放进度手动控制"),
            ("Low", "3", "资金保障、批量分配、支付代币兼容性等运营/边界风险"),
            ("Info", "2", "升级预留、上线配置说明"),
        ],
        widths=[1.2, 0.7, 4.6],
    )

    add_heading(doc, "2. 审计范围", 1)
    add_table(
        doc,
        ["类别", "文件/对象", "状态"],
        [
            ("Token", "contracts/PESToken.sol", "审计"),
            ("Token UUPS", "contracts/PESTokenUpgradeable.sol", "审计 / 当前主网使用"),
            ("Presale", "contracts/PESPresaleVesting.sol", "审计"),
            ("Presale UUPS", "contracts/PESPresaleVestingUpgradeable.sol", "审计 / 当前主网使用"),
            ("Tests", "test/PESToken.test.js; test/PESPresaleVesting.test.js; test/Upgradeable.test.js", "验证"),
            ("Deployments", "deployments/bsc-mainnet-upgradeable-2026-05-24.json; deployments/bsc-mainnet-presale-upgrade-2026-05-25.json", "核对"),
        ],
        widths=[1.2, 4.1, 1.2],
    )
    add_para(doc, "不在本次审计范围内：BSC USDT 第三方合约、钱包私钥管理、Cloudflare 前端/API 安全、商业定价合理性、中心化交易所/DEX 流动性安排、形式化验证。")

    add_heading(doc, "3. 主网状态快照", 1)
    add_para(doc, "只读链上核验时间：2026-05-26T03:03:49Z，BSC block 100464401。")
    add_table(
        doc,
        ["项目", "当前值"],
        [
            ("PES Token Proxy", "0xe83e750feEbe231c870DdF30165CbFE64F400Ebc"),
            ("PES Token Implementation", "0x7909b4c672059D00765c0963A3ba4f2219E90B82"),
            ("Presale Vesting Proxy", "0x6d5Fc8F6A0481a81A726Ca2Fac85c23ED80619fd"),
            ("Presale Vesting Implementation", "0x57B67860ca1941b4fC1EFb98D58c693639183A8a"),
            ("Payment Token", "USDT 0x55d398326f99059fF775485246999027B3197955"),
            ("Owner", "0xAC25dA7FdEEEaDf2943EBF505Fa9739CBD111bD8"),
            ("Funds Wallet", "0x69D713beEAB0e75e0c71C5A4b054b46A83E4c6c1"),
            ("PES Total Supply", "21,000,000 PES"),
            ("Presale PES Balance", "3,000,000 PES"),
            ("Payment / Package", "300 USDT"),
            ("PES / Package", "3,000 PES"),
            ("Max/Public Cap", "1,000 / 1,000 packages"),
            ("Per Wallet Limit", "1 package"),
            ("Public Sold / Allocated", "1 / 1 package"),
            ("Trading Enabled", "false"),
            ("Presale Paused", "false"),
            ("Launch Time", "0（未设置）"),
            ("Elapsed Vesting Periods", "0"),
        ],
        widths=[2.2, 4.2],
    )

    add_heading(doc, "4. 正向安全观察", 1)
    for item in [
        "使用 OpenZeppelin ERC20、Ownable、Pausable、SafeERC20 与 UUPS 组件，基础依赖来源清晰。",
        "Token 买卖费率总上限为 10%，当前买卖费率均为 150 bps。",
        "预售购买、领取使用 nonReentrant，领取先更新状态再转账，重入风险被控制。",
        "预售总份数、公售份数、单钱包份数均有链上限制。",
        "recoverUnsupportedToken 对已分配但未领取的 PES 做 reserved 保护。",
        "测试覆盖购买、限购、Admin 分配、释放进度、未充值先购买、充值后领取、UUPS 升级路径。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. 审计发现", 1)
    add_finding(
        doc,
        "DWP-PES-M01",
        "Medium",
        "单一 Owner 控制升级、暂停、费率、释放和销售参数",
        "PESTokenUpgradeable.sol:80-150 暴露 pause/unpause、setTradingEnabled、setAutomatedMarketMakerPair、setFeeWallets、setFeeRates、_authorizeUpgrade；PESPresaleVestingUpgradeable.sol:134-176 和 309 暴露预售参数、资金钱包、释放配置和升级权限。主网 owner 为 0xAC25...1bD8。",
        "如果 owner 私钥被盗或误操作，攻击者可升级实现、暂停系统、调整费率/交易对/资金钱包、变更预售窗口或推进释放，影响用户资产和上线可信度。",
        "将 owner 迁移到多签钱包，生产升级引入 timelock 或至少 24 小时公告窗口；将费率、升级、释放、资金钱包操作拆分为不同角色；每次参数变更保留链上/后台操作记录。",
    )
    add_finding(
        doc,
        "DWP-PES-M02",
        "Medium",
        "释放进度由 Owner 手动推进，launchTime/vestingPeriodSeconds 未直接约束 claim",
        "PESPresaleVestingUpgradeable.sol:252-270 的 vestedAmount 只读取 elapsedVestingPeriods；159-169 允许 owner 设置释放进度；344-356 只限制不能回退、不能超过 vestingPeriods + 1。launchTime 和 vestingPeriodSeconds 当前不参与 claimableAmount 计算。",
        "释放可以被 owner 提前推进到全额，也可能因 owner 未推进导致用户无法按预期领取。该设计更像“运营手动释放”，而不是完全链上时间锁释放。",
        "如果业务目标是自动释放，应按 launchTime + vestingPeriodSeconds 计算 elapsed periods，并只允许 owner 修正异常；如果保留手动释放，应在白皮书/页面明确说明，并建立每日释放操作和复核流程。",
    )
    add_finding(
        doc,
        "DWP-PES-L01",
        "Low",
        "预售允许先分配后充值，领取前资金保障依赖运营检查",
        "PESPresaleVestingUpgradeable.sol:196-221 和 223-236 不检查合约 PES 余额；claim 时通过 SafeERC20 转账失败来阻止领取。当前主网余额为 3,000,000 PES，覆盖 1,000 份上限。",
        "当前主网状态已满足资金覆盖，但未来如果提高 maxPackages、修改 pesPerPackage 或恢复/回收代币，可能出现用户已付款但领取失败。",
        "保留上线前/每次参数变更后的 funding checklist：required = maxPackages * pesPerPackage；presalePesBalance 必须大于等于 unclaimedAllocatedTokens 和最大销售承诺。前端和 Admin 面板建议显示释放资金覆盖率。",
    )
    add_finding(
        doc,
        "DWP-PES-L02",
        "Low",
        "grantAllocations 批量分配没有链上批次上限",
        "PESPresaleVestingUpgradeable.sol:228-236 对 accounts 数组逐项循环，没有最大长度限制。",
        "过大的批量输入会因为 gas 超限失败；虽然不会造成资金损失，但会影响运营批量发放效率，并增加误操作概率。",
        "在合约层加入 MAX_BATCH_SIZE，或在 Admin 工具强制分片并展示每批 gas 估算。当前前端已有 chunk size 概念，应在操作手册中固定推荐批次。",
    )
    add_finding(
        doc,
        "DWP-PES-L03",
        "Low",
        "支付代币按名义金额转账，未校验 fundsWallet 实收差额",
        "purchasePackages 在 PESPresaleVestingUpgradeable.sol:215-218 计算 paymentAmount 后调用 paymentToken.safeTransferFrom(msg.sender, fundsWallet, paymentAmount)，未检查 fundsWallet 转账前后余额差。",
        "当前主网支付代币为 BSC USDT，通常可接受；但如果未来替换为 fee-on-transfer 或异常 ERC20，可能出现实收低于名义价格。",
        "主网保持 USDT 地址不可变；如未来支持可更换支付代币，应增加 balance delta 校验或明确只接受标准 ERC20。",
    )
    add_finding(
        doc,
        "DWP-PES-I01",
        "Info",
        "升级合约未预留 storage gap",
        "PESTokenUpgradeable.sol 与 PESPresaleVestingUpgradeable.sol 使用 UUPS，并依赖 OpenZeppelin storage validation，但合约末尾未预留 __gap。",
        "短期不构成漏洞；后续升级若需要插入继承合约或重排状态变量，改动空间较小。",
        "未来版本只在末尾追加变量，并保留 OpenZeppelin validateUpgrade；可在下一次兼容升级中追加 storage gap。",
        status="Advisory",
    )
    add_finding(
        doc,
        "DWP-PES-I02",
        "Info",
        "上线参数仍需要运营公告和冻结策略",
        "链上当前 tradingEnabled=false，launchTime=0，elapsedVestingPeriods=0；saleStart 已开始，saleEnd 为 2026-07-25T08:36:00Z。",
        "用户可以认购，但交易开放和领取释放仍依赖后续操作。缺少清晰公告会造成用户误解。",
        "正式开售公告中列明：认购窗口、单份价格、领取开始时间、释放规则、预售合约地址、PES 资金覆盖状态、Owner/多签地址。",
        status="Advisory",
    )

    add_heading(doc, "6. 测试与验证", 1)
    add_table(
        doc,
        ["检查项", "结果"],
        [
            ("npm test", "16 passing"),
            ("npm run compile", "Nothing to compile / success"),
            ("UUPS upgrade smoke test", "通过 Upgradeable.test.js"),
            ("预售未充值先购买测试", "通过"),
            ("释放进度不可回退/不可越界测试", "通过"),
            ("主网只读核验", "BSC chainId 56; block 100464401"),
        ],
        widths=[2.2, 4.1],
    )

    add_heading(doc, "7. 上线前检查清单", 1)
    for item in [
        "确认 owner 已迁移到多签，或至少明确单签临时控制期限。",
        "确认 PES 预售合约余额 >= maxPackages * pesPerPackage；当前为 3,000,000 PES。",
        "确认 paymentPerPackage = 300 USDT，perWalletPackageLimit = 1，publicPackageCap = 1,000。",
        "设置并公告 launchTime，或明确释放由 Admin 手动推进。",
        "正式开放交易前确认 tradingEnabled、AMM pair、fee wallet、buy/sell fee bps。",
        "每次释放前读取 elapsedVestingPeriods、claimableAmount 抽样账户、presalePesBalance。",
        "保留升级前 storage validation、implementation 地址、tx hash 和回滚预案。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. 附录：DAPPWEB 品牌和方法", 1)
    add_para(doc, "DAPPWEB 品牌定位为 Web3 安全与产品工程伙伴，安全交付重点覆盖 scope lock、threat model、manual review、tests、patch/retest 和 launch checklist。")
    add_para(doc, "品牌参考来源：https://dappwebofficialwebsite.pages.dev/")
    add_para(doc, "报告生成路径：output/DAPPWEB-PES-Contract-Audit-Report-2026-05-26.docx")

    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.text = "DAPPWEB · PES Contract Security Audit"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            set_run_font(run, size=8, color=COLORS["muted"])
        fp = sec.footer.paragraphs[0]
        fp.text = "Confidential Review Draft · 2026-05-26"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            set_run_font(run, size=8, color=COLORS["muted"])

    doc.core_properties.title = "DAPPWEB PES Contract Audit Report"
    doc.core_properties.subject = "PES Token and Presale Vesting smart contract audit"
    doc.core_properties.author = "DAPPWEB"
    doc.core_properties.keywords = "DAPPWEB, PES, Smart Contract Audit, BSC, UUPS, Presale, Vesting"

    OUTPUT.parent.mkdir(exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT.resolve())

"""Generate user distribution audit report (DOCX) from audit JSON."""
from __future__ import annotations

import json
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output"


def latest_audit_json() -> Path:
    files = sorted(OUTPUT_DIR.glob("pes-user-distribution-audit-*.json"))
    if not files:
        raise FileNotFoundError("Run node output/audit-user-distributions.js first")
    return files[-1]


def rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.strip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_run_font(run, size=10, color=None, bold=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "0B1020")
        set_cell_text(cell, header, bold=True, color="FFFFFF", size=9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def add_heading(doc, text, level=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=13 if level == 2 else 11, bold=True, color="0B1020")
    p.paragraph_format.space_after = Pt(6)


def add_para(doc, text, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, color="6B7280")
    p.paragraph_format.space_after = Pt(4)


def claim_histogram(accounts: list) -> dict[str, int]:
    hist: dict[str, int] = {}
    for row in accounts:
        key = row["claimedPES"]
        hist[key] = hist.get(key, 0) + 1
    return dict(sorted(hist.items(), key=lambda x: -float(x[0])))


def build_doc(data: dict, output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PES 已发放用户审计报告")
    set_run_font(run, size=18, bold=True, color="0B1020")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"BSC 主网 · 区块 #{data['blockNumber']} · {data['blockTimeUtc'][:19].replace('T', ' ')} UTC"
    )
    set_run_font(run, size=10, color="6B7280")

    s = data["summary"]
    v = data["vesting"]
    contracts = data["contracts"]

    add_para(doc, f"归属合约：{contracts['presaleVesting']}")
    add_para(doc, f"PES 代币：{contracts['pesToken']}")

    add_heading(doc, "一、审计结论", 2)
    conclusions = [
        f"受益用户数：{s['beneficiaryCount']} 人，归属登记总量 {s['totalAllocatedPES']} PES。",
        f"链上 totalTokensClaimed = {s['totalClaimedOnChainPES']} PES，与各地址 allocation.claimed 合计一致：{'是' if s['onChainMatchesAllocationSum'] else '否'}。",
        f"当前已执行第 {v['elapsedPeriods']} 期，释放比例 {v['releasePercent']}，每人应释放 {v['perUserExpectedPES']} PES。",
        f"已全额发放至应发额度：{s['usersFullyAtExpected']} 人；部分发放（未达当期应发）：{s['usersPartialClaim']} 人；超额发放：{s['usersOverVested']} 人。",
        "部分发放原因为归属合约 PES 余额不足，发放批次在约第 321 名用户之后中断。",
    ]
    for line in conclusions:
        add_para(doc, f"• {line}", size=10)

    add_heading(doc, "二、发放汇总", 2)
    add_table(
        doc,
        ["指标", "数值"],
        [
            ["受益用户数", str(s["beneficiaryCount"])],
            ["归属登记总量 (PES)", s["totalAllocatedPES"]],
            ["已发放到用户钱包 (PES)", s["totalClaimedOnChainPES"]],
            ["当期应释放合计 (PES)", s["sumVestedPES"]],
            ["未发放差额 (应发−已发)", f"{float(s['sumVestedPES']) - float(s['totalClaimedOnChainPES']):,.0f}"],
            ["已执行期数", str(v["elapsedPeriods"])],
            ["当期释放比例", v["releasePercent"]],
            ["每人当期应发 (PES)", v["perUserExpectedPES"]],
        ],
        widths=[2.8, 2.8],
    )

    add_heading(doc, "三、已领取金额分布", 2)
    hist = claim_histogram(data["accounts"])
    add_table(
        doc,
        ["已领取 (PES/人)", "用户数", "合计 (PES)"],
        [
            [amount, str(count), f"{float(amount) * count:,.0f}"]
            for amount, count in hist.items()
        ],
        widths=[1.8, 1.2, 1.8],
    )

    add_heading(doc, "四、发放批次记录", 2)
    batches = data.get("distributionBatches", [])
    batch_rows = []
    for batch in batches:
        if "txHash" in batch:
            batch_rows.append([
                str(batch.get("elapsedPeriods", batch.get("period", ""))),
                str(batch.get("accountCount", batch.get("accounts", ""))),
                batch.get("distributedPES", ""),
                str(batch.get("blockNumber", "")),
            ])
        else:
            batch_rows.append([
                str(batch.get("period", "")),
                str(batch.get("accounts", "")),
                batch.get("distributedPES", ""),
                batch.get("note", ""),
            ])
    add_table(
        doc,
        ["期数", "账户数", "发放量 (PES)", "备注/区块"],
        batch_rows or [["—", "—", "—", "—"]],
        widths=[1.0, 1.0, 1.4, 2.2],
    )

    add_heading(doc, "五、异常账户（部分发放）", 2)
    anomalies = data.get("anomalies", [])
    add_para(doc, f"共 {len(anomalies)} 个账户未达当期应发额度（示例前 20 条，完整清单见 CSV/JSON）。")
    sample = anomalies[:20]
    add_table(
        doc,
        ["序号", "地址", "已领取", "应发", "差额", "首次购买时间"],
        [
            [
                str(a["index"]),
                a["account"],
                a["claimedPES"],
                a["expectedAtElapsedPES"],
                f"{float(a['expectedAtElapsedPES']) - float(a['claimedPES']):.0f}",
                a.get("firstPurchaseTime", ""),
            ]
            for a in sample
        ],
        widths=[0.5, 2.2, 0.8, 0.8, 0.6, 1.5],
    )

    add_heading(doc, "六、审计方法", 2)
    methods = [
        "数据源：当前归属合约 allocation 映射（packages / tokens / claimed），不含旧合约。",
        "逐地址读取 1000 名历史买家的链上 claimed 字段，与 totalTokensClaimed 交叉核对。",
        "应发金额按 elapsedVestingPeriods 与释放公式计算（首笔 20%，之后每期间线性 +2%）。",
        "status=ok 表示 claimed 等于当期应发；status=partial 表示 claimed 低于当期应发。",
        f"明细文件：{OUTPUT_DIR / data.get('csvPath', 'pes-user-distribution-audit-*.csv').split('/')[-1]}",
    ]
    for line in methods:
        add_para(doc, f"• {line}")

    add_para(doc, f"报告生成：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')} UTC")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def convert_pdf(docx_path: Path) -> Path | None:
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        return pdf_path
    except Exception as error:
        print(f"PDF 转换跳过: {error}", file=sys.stderr)
        return None


def main() -> None:
    audit_path = Path(sys.argv[1]) if len(sys.argv) > 1 else latest_audit_json()
    data = json.loads(audit_path.read_text(encoding="utf-8"))
    stamp = audit_path.stem.replace("pes-user-distribution-audit-", "")
    docx_path = OUTPUT_DIR / f"pes-user-distribution-audit-{stamp}.docx"

    build_doc(data, docx_path)
    print(str(docx_path))

    pdf_path = convert_pdf(docx_path)
    if pdf_path:
        print(str(pdf_path))


if __name__ == "__main__":
    main()
